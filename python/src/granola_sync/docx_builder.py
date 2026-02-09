"""Build styled .docx files from meeting data."""

from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .html_parser import parse_html_to_elements


def _add_hr(doc: Document) -> None:
    """Add a light gray horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): "CCCCCC",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_meeting_docx(
    filepath: str,
    title: str,
    date_str: str,
    attendees: list[dict],
    summary_html: str,
    transcript_chunks: list[dict],
    notes_markdown: str | None = None,
) -> None:
    doc = Document()

    # Global style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # Title
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Date
    try:
        dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        formatted_date = dt.strftime("%A, %B %d, %Y  |  %I:%M %p")
    except (ValueError, AttributeError):
        formatted_date = date_str
    p = doc.add_paragraph()
    run = p.add_run(formatted_date)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    # Attendees
    if attendees:
        names = [a.get("name", a.get("email", "Unknown")) for a in attendees]
        p = doc.add_paragraph()
        run = p.add_run("Attendees:  ")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(", ".join(names)).font.size = Pt(11)

    _add_hr(doc)

    # Summary
    if summary_html:
        h = doc.add_heading("Meeting Summary", level=2)
        for run in h.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        elements = parse_html_to_elements(summary_html)
        for elem_type, text, depth in elements:
            if elem_type == "heading":
                sh = doc.add_heading(text, level=3)
                for run in sh.runs:
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            elif elem_type == "bullet":
                p = doc.add_paragraph(style="List Bullet")
                p.text = text
                if depth > 1:
                    p.paragraph_format.left_indent = Inches(0.25 * depth)
            elif elem_type == "paragraph":
                doc.add_paragraph(text)

    # Notes
    if notes_markdown and notes_markdown.strip():
        _add_hr(doc)
        h = doc.add_heading("Notes", level=2)
        for run in h.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        for line in notes_markdown.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith(("- ", "* ")):
                p = doc.add_paragraph(style="List Bullet")
                p.text = stripped[2:]
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=4)
            else:
                doc.add_paragraph(stripped)

    # Transcript
    if transcript_chunks:
        doc.add_page_break()
        h = doc.add_heading("Full Transcript", level=2)
        for run in h.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        for chunk in transcript_chunks:
            ts = chunk.get("start_timestamp", "")
            source = chunk.get("source", "unknown")
            text = chunk.get("text", "").strip()
            if not text:
                continue

            try:
                cdt = datetime.fromisoformat(ts.replace("Z", "+00:00"))
                time_str = cdt.strftime("%H:%M:%S")
            except (ValueError, AttributeError):
                time_str = "??:??:??"

            source_label = "\U0001f3a4 You" if source == "microphone" else "\U0001f50a Other"

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)

            ts_run = p.add_run(f"{time_str}  ")
            ts_run.font.size = Pt(8)
            ts_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            label_run = p.add_run(f"[{source_label}]  ")
            label_run.font.size = Pt(9)
            label_run.bold = True
            if source == "microphone":
                label_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
            else:
                label_run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

            text_run = p.add_run(text)
            text_run.font.size = Pt(9.5)

    doc.save(filepath)
